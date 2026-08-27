import io
from docx import Document
from docx.shared import Pt
from database import INDICATORS, INDICATOR_KEYS


async def generate_report_docx(summary_rows: list, detailed_rows: list = None) -> io.BytesIO:
    """
    Генерирует официальный Word-отчёт АДТИ по форме «2026 йил хисобот»
    1-қисм: 65 та кафедра сводкаси
    2-қисм: Илмий ишлар ва уларнинг муаллифлари батафсил рўйхати (Шаффофлик учун)
    """
    from database import INDICATOR_LABELS

    doc = Document()

    # ─── 1-ҚИСМ: СВОДКА ТАБЛИЦАСИ ────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = 1  # CENTER
    run = title.add_run("АНДИЖОН ДАВЛАТ ТИББИЁТ ИНСТИТУТИ КАФЕДРАЛАРИ ТОМОНИДАН\n")
    run.bold = True
    run.font.size = Pt(12)
    run2 = title.add_run("2026 ЙИЛ ХИСОБОТИ — ИЛМИЙ ТАДҚИҚОТ ИШЛАРИ ТЎҒРИСИДА МАЪЛУМОТ")
    run2.bold = True
    run2.font.size = Pt(12)

    doc.add_paragraph()

    # Таблица: 17 столбцов как в оригинале
    col_headers = [
        "Т/Р", "Кафедра номи", "Кафедра мудири",
        "DSc", "PhD", "Монография", "Патент",
        "ЎзОАК", "Россия ОАК/IF",
        "Тезис (Ўз)", "Тезис (Хор)",
        "Scopus/WoS",
        "Рац.таклиф", "Амалиётга тадбиқ",
        "Анжуман", "Хўж.шартн.", "Грант"
    ]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = "Table Grid"

    # Заголовки
    hdr = table.rows[0].cells
    for i, h in enumerate(col_headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    # Данные
    totals = [0] * (len(col_headers) - 3)  # для итоговой строки

    for r in summary_rows:
        row = table.add_row().cells
        row[0].text = str(r['id'])
        row[1].text = r['name']
        row[2].text = r['head_name'] or ''

        vals = [
            r['dsc'], r['phd'], r['monography'], r['patent'],
            r['oak_uz'], r['oak_ru_if'],
            r['thesis_uz'], r['thesis_foreign'],
            r['scopus_wos'],
            r['rationalizer'], r['implementation'],
            r['conferences'], '', ''  # contracts и grants — текстовые поля
        ]

        for i, v in enumerate(vals):
            row[i + 3].text = str(v) if v else ''
            if isinstance(v, int):
                totals[i] += v

    # Итоговая строка
    total_row = table.add_row().cells
    total_row[0].text = ''
    total_row[1].text = 'ЖАМИ (Итого):'
    total_row[2].text = ''
    for i, t in enumerate(totals):
        total_row[i + 3].text = str(t) if t else ''

    # Формат шрифтов таблицы 1
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # ─── 2-ҚИСМ: МУАЛЛИФЛАР ВА ИШЛАР БАТАФСИЛ РЎЙХАТИ (ИЛОВА) ─────────────────
    if detailed_rows:
        doc.add_page_break()

        p2 = doc.add_paragraph()
        p2.alignment = 1
        r_app = p2.add_run("ИЛОВА — 2-ҚИСМ\n")
        r_app.bold = True
        r_app.font.size = Pt(11)
        r_app2 = p2.add_run("КАФЕДРАЛАР КЕСИМИДА ИЛМИЙ ИШЛАР ВА УЛАРНИНГ МУАЛЛИФЛАРИ БАТАФСИЛ РЎЙХАТИ\n(ШАФФОФЛИК ҲИСОБОТИ)")
        r_app2.bold = True
        r_app2.font.size = Pt(12)

        p_info = doc.add_paragraph()
        p_info.add_run(f"Жами рўйхатга олинган ишлар: {len(detailed_rows)} та\n").italic = True

        det_cols = [
            "Т/Р", "Кафедра номи", "Йўналиш",
            "Иш номи / Мавзу", "Муаллифлар (Ф.И.Ш.)", "Қўшимча маълумот (Журнал/Сана)"
        ]

        det_table = doc.add_table(rows=1, cols=len(det_cols))
        det_table.style = "Table Grid"

        det_hdr = det_table.rows[0].cells
        for i, h in enumerate(det_cols):
            det_hdr[i].text = h
            det_hdr[i].paragraphs[0].runs[0].bold = True

        for idx, d in enumerate(detailed_rows, 1):
            row = det_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = str(d.get('dept_name', '') or f"Кафедра #{d.get('dept_id', '')}")
            cat_k = d.get('category', '')
            row[2].text = INDICATOR_LABELS.get(cat_k, cat_k)
            row[3].text = str(d.get('title', '') or '—')
            
            # Муаллифлар алоҳида ажралиб туради (Bold)
            row[4].text = str(d.get('authors', '') or '—')
            if row[4].paragraphs and row[4].paragraphs[0].runs:
                row[4].paragraphs[0].runs[0].bold = True

            # Қўшимча майдонлар (журнал, сана, давлат, сумма)
            extra_parts = []
            if d.get('country'): extra_parts.append(f"Давлат: {d['country']}")
            if d.get('journal_name'): extra_parts.append(f"Журнал: {d['journal_name']}")
            if d.get('pub_date'): extra_parts.append(f"Сана: {d['pub_date']}")
            if d.get('amount'): extra_parts.append(f"Сумма: {d['amount']} млн")
            if d.get('reg_number'): extra_parts.append(f"Рег: {d['reg_number']}")
            row[5].text = "; ".join(extra_parts) if extra_parts else "—"

        for row in det_table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


async def generate_codes_docx(departments: list) -> io.BytesIO:
    """
    Генерирует Word-файл с индивидуальными кодами доступа для всех 65 кафедр
    (для передачи завкафедрами)
    """
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("АНДИЖОН ДАВЛАТ ТИББИЁТ ИНСТИТУТИ\n")
    run.bold = True
    run.font.size = Pt(13)
    run2 = title.add_run("КАФЕДРАЛАР УЧУН ТЕЛЕГРАМ БОТГА КИРИШ МАХСУС ПАРОЛЛАРИ (КОДЛАРИ)")
    run2.bold = True
    run2.font.size = Pt(11)

    doc.add_paragraph("Ушбу махсус кодлар ҳар бир кафедра мудири ёки масъул ходимига берилади. Ботга биринчи марта кирганда ушбу код киритилади.\n")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "Т/Р"
    hdr[1].text = "Кафедра номи"
    hdr[2].text = "Кафедра мудири"
    hdr[3].text = "КИРИШ КОДИ (ПАРОЛ)"

    for i in range(4):
        hdr[i].paragraphs[0].runs[0].bold = True

    for dep_id, name, head, code in departments:
        row = table.add_row().cells
        row[0].text = str(dep_id)
        row[1].text = name
        row[2].text = head or ''
        row[3].text = code
        row[3].paragraphs[0].runs[0].bold = True

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


async def generate_report_excel(summary_rows: list, detailed_rows: list) -> io.BytesIO:
    """
    Генерирует официальный Excel (.xlsx) отчёт с двумя листами:
    1. «Сводка» — таблица 65 кафедр по всем 17 индикаторам с авто-суммами
    2. «Барча ҳисоботлар» — полная база всех отправленных записей
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # Стили
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    total_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
    bold_font = Font(name="Calibri", size=10, bold=True)
    regular_font = Font(name="Calibri", size=10)
    
    thin_border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9")
    )
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # ─── ЛИСТ 1: СВОДКА ──────────────────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Сводка 65 кафедр"
    ws1.views.sheetView[0].showGridLines = True

    # Заголовок
    ws1.merge_cells("A1:R1")
    title_cell = ws1["A1"]
    title_cell.value = "АНДИЖОН ДАВЛАТ ТИББИЁТ ИНСТИТУТИ — 2026 ЙИЛ ИЛМИЙ ХИСОБОТИ"
    title_cell.font = Font(name="Calibri", size=14, bold=True, color="1F4E79")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 30

    col_headers = [
        "Т/Р", "Кафедра номи", "Кафедра мудири",
        "DSc", "PhD", "Монография", "Патент",
        "ЎзОАК", "Россия ОАК/IF",
        "Тезис (Ўз)", "Тезис (Хор)",
        "Scopus/WoS",
        "Рац.таклиф", "Амалиётга тадбиқ",
        "Анжуман", "Хўж.шартн.", "Грант", "Жами"
    ]

    ws1.append([])  # Строка 2 пустая
    ws1.append(col_headers)  # Строка 3 заголовки
    ws1.row_dimensions[3].height = 28

    for col_num in range(1, len(col_headers) + 1):
        c = ws1.cell(row=3, column=col_num)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center_align
        c.border = thin_border

    totals = [0] * (len(col_headers) - 3)

    for r in summary_rows:
        row_vals = [
            r['id'], r['name'], r['head_name'] or '',
            r['dsc'] or 0, r['phd'] or 0, r['monography'] or 0, r['patent'] or 0,
            r['oak_uz'] or 0, r['oak_ru_if'] or 0,
            r['thesis_uz'] or 0, r['thesis_foreign'] or 0,
            r['scopus_wos'] or 0,
            r['rationalizer'] or 0, r['implementation'] or 0,
            r['conferences'] or 0, r.get('contracts', 0) or 0, r.get('grants', 0) or 0,
            r['total'] or 0
        ]
        ws1.append(row_vals)
        curr_row = ws1.max_row
        ws1.row_dimensions[curr_row].height = 20

        for idx, val in enumerate(row_vals):
            cell = ws1.cell(row=curr_row, column=idx + 1)
            cell.font = regular_font
            cell.border = thin_border
            if idx == 1:
                cell.alignment = left_align
            else:
                cell.alignment = center_align

            if idx >= 3 and isinstance(val, (int, float)):
                totals[idx - 3] += val

    # Итоговая строка
    total_row_vals = ["", "ЖАМИ (Институт бўйича):", ""] + totals
    ws1.append(total_row_vals)
    tot_row_num = ws1.max_row
    ws1.row_dimensions[tot_row_num].height = 24

    for col_num in range(1, len(total_row_vals) + 1):
        cell = ws1.cell(row=tot_row_num, column=col_num)
        cell.font = bold_font
        cell.fill = total_fill
        cell.border = thin_border
        cell.alignment = left_align if col_num == 2 else center_align

    # ─── ЛИСТ 2: ДЕТАЛЬНЫЕ ЗАПИСИ ────────────────────────────────────────────
    ws2 = wb.create_sheet(title="Барча ҳисоботлар (База)")
    ws2.views.sheetView[0].showGridLines = True

    from database import INDICATOR_LABELS
    det_headers = [
        "ID", "Сана/Вақт", "Кафедра ID", "Кафедра номи", "Кафедра мудири",
        "Категория", "Иш номи / Диссертация мавзуси", "Муаллифлар (Ф.И.Ш.)",
        "Шартнома/Грант суммаси (млн)", "Нашр давлати", "Журнал/Буюртмачи", "Нашр йили / Сана / Бетлар",
        "URL / DOI", "Муаллифлар сони", "Ихтисослик шифри ва номи",
        "Рег. рақам (Патент)", "Нашриёт (Монография)", "Файл борми?"
    ]
    ws2.append(det_headers)
    ws2.row_dimensions[1].height = 26

    for col_num in range(1, len(det_headers) + 1):
        c = ws2.cell(row=1, column=col_num)
        c.font = header_font
        c.fill = header_fill
        c.alignment = center_align
        c.border = thin_border

    for d in detailed_rows:
        has_f = "✅ Ҳа" if d['file_path'] else "❌ Йўқ"
        cat_lbl = INDICATOR_LABELS.get(d['category'], d['category'])
        doi_or_url = d.get('doi', '') or d.get('url', '') or ''
        r_data = [
            d['id'], str(d['created_at'])[:16], d['dept_id'], d['dept_name'],
            d['head_name'] or '', cat_lbl,
            d['title'] or '',
            d['authors'] or '',
            d.get('amount', '') or '',
            d.get('country', '') or '',
            d.get('journal_name', '') or '',
            d.get('pub_date', '') or '',
            doi_or_url,
            d.get('authors_count', '') or '',
            d.get('specialty', '') or '',
            d.get('reg_number', '') or '',
            d.get('publisher', '') or '',
            has_f
        ]
        ws2.append(r_data)
        curr = ws2.max_row
        for col_idx in range(1, len(r_data) + 1):
            cell = ws2.cell(row=curr, column=col_idx)
            cell.font = regular_font
            cell.border = thin_border
            cell.alignment = left_align if col_idx in (4, 7, 8, 9, 10, 11, 12, 13, 15) else center_align

    # ─── ЛИСТ 3: КАФЕДРАЛАР РЕЙТИНГИ ВА ЛИДЕРЛАР ────────────────────────────
    ws3 = wb.create_sheet(title="Кафедралар рейтинги")
    ws3.views.sheetView[0].showGridLines = True

    # Заголовок Листа 3
    ws3.merge_cells("A1:N1")
    title_r = ws3["A1"]
    title_r.value = "АНДИЖОН ДАВЛАТ ТИББИЁТ ИНСТИТУТИ — КАФЕДРАЛАРНИНГ ИЛМИЙ ФАОЛЛИК РЕЙТИНГИ (2026 ЙИЛ)"
    title_r.font = Font(name="Calibri", size=13, bold=True, color="1F4E79")
    title_r.alignment = Alignment(horizontal="center", vertical="center")
    ws3.row_dimensions[1].height = 28

    rank_headers = [
        "Ўрни (Рейтинг)", "Кафедра ID", "Кафедра номи", "Кафедра мудири",
        "ЖАМИ ИШЛАР", "Scopus / WoS", "ЎзОАК мақола", "Россия ОАК / IF",
        "Патентлар", "DSc / PhD", "Монография", "Тезислар (Ўз/Хор)",
        "Шартнома / Грант", "Фаоллик ҳолати"
    ]
    ws3.append([])
    ws3.append(rank_headers)
    ws3.row_dimensions[3].height = 26

    rank_header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    top10_fill = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")  # soft green
    zero_fill = PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid")   # soft red

    for col_num in range(1, len(rank_headers) + 1):
        c = ws3.cell(row=3, column=col_num)
        c.font = header_font
        c.fill = rank_header_fill
        c.alignment = center_align
        c.border = thin_border

    # Сортировка кафедр по убыванию количества работ
    sorted_depts = sorted(summary_rows, key=lambda r: -(r.get('total') or 0))

    for rank, r in enumerate(sorted_depts, 1):
        tot = r.get('total') or 0
        diss = (r.get('dsc') or 0) + (r.get('phd') or 0)
        theses = (r.get('thesis_uz') or 0) + (r.get('thesis_foreign') or 0)
        grants = (r.get('contracts') or 0) + (r.get('grants') or 0)

        if tot >= 50:
            status_txt = "🔥 Юқори фаол (Лидер)"
        elif tot >= 15:
            status_txt = "✅ Фаол"
        elif tot > 0:
            status_txt = "⚠️ Паст кўрсаткич"
        else:
            status_txt = "❌ Иш топширмаган (0)"

        row_vals = [
            f"#{rank}", r['id'], r['name'], r['head_name'] or '—',
            tot, r['scopus_wos'] or 0, r['oak_uz'] or 0, r['oak_ru_if'] or 0,
            r['patent'] or 0, diss, r['monography'] or 0, theses,
            grants, status_txt
        ]
        ws3.append(row_vals)
        curr = ws3.max_row
        ws3.row_dimensions[curr].height = 20

        is_top10 = rank <= 10 and tot > 0
        is_zero = tot == 0

        for col_idx in range(1, len(row_vals) + 1):
            cell = ws3.cell(row=curr, column=col_idx)
            cell.font = bold_font if col_idx in (1, 5, 14) else regular_font
            cell.border = thin_border
            if is_top10:
                cell.fill = top10_fill
            elif is_zero:
                cell.fill = zero_fill

            if col_idx in (3, 4):
                cell.alignment = left_align
            else:
                cell.alignment = center_align

    # ─── ТАБЛИЦА 2: ТОП-30 МУАЛЛИФЛАР (ЛИСТ 3) ──────────────────────────────
    ws3.append([])
    ws3.append([])
    author_table_start = ws3.max_row + 1

    ws3.merge_cells(f"A{author_table_start}:H{author_table_start}")
    auth_title = ws3[f"A{author_table_start}"]
    auth_title.value = "ИНСТИТУТНИНГ ЭНГ ФАОЛ МУАЛЛИФЛАРИ ВА ТАДҚИҚОТЧИЛАРИ (ТОП-30)"
    auth_title.font = Font(name="Calibri", size=12, bold=True, color="1F4E79")
    auth_title.alignment = Alignment(horizontal="center", vertical="center")
    ws3.row_dimensions[author_table_start].height = 26

    auth_headers = [
        "Ўрни", "Муаллифнинг Ф.И.Ш.", "Кафедра номи",
        "ЖАМИ ИШЛАРИ", "Scopus / WoS", "Патентлар", "ЎзОАК / ОАК", "Диссертация / Тезислар"
    ]
    ws3.append(auth_headers)
    hdr_row_idx = ws3.max_row
    ws3.row_dimensions[hdr_row_idx].height = 24
    for col_num in range(1, len(auth_headers) + 1):
        c = ws3.cell(row=hdr_row_idx, column=col_num)
        c.font = header_font
        c.fill = rank_header_fill
        c.alignment = center_align
        c.border = thin_border

    # Подсчёт активности авторов
    authors_data = {}
    for d in detailed_rows:
        auth_str = (d.get('authors') or '').strip()
        dept_name = d.get('dept_name') or ''
        cat = d.get('category') or ''
        for raw_a in auth_str.replace(";", ",").split(","):
            a = raw_a.strip()
            if a and len(a) > 2 and a.lower() not in ["va boshqalar", "et al", "—"]:
                if a not in authors_data:
                    authors_data[a] = {"dept": dept_name, "total": 0, "scopus": 0, "patent": 0, "oak": 0, "other": 0}
                authors_data[a]["total"] += 1
                if not authors_data[a]["dept"] and dept_name:
                    authors_data[a]["dept"] = dept_name
                if cat == "scopus_wos":
                    authors_data[a]["scopus"] += 1
                elif cat == "patent":
                    authors_data[a]["patent"] += 1
                elif cat in ("oak_uz", "oak_ru_if"):
                    authors_data[a]["oak"] += 1
                else:
                    authors_data[a]["other"] += 1

    sorted_authors = sorted(authors_data.items(), key=lambda x: -x[1]["total"])

    for a_rank, (a_name, a_info) in enumerate(sorted_authors[:30], 1):
        a_row = [
            f"#{a_rank}", a_name, a_info["dept"] or "—",
            a_info["total"], a_info["scopus"], a_info["patent"],
            a_info["oak"], a_info["other"]
        ]
        ws3.append(a_row)
        curr = ws3.max_row
        ws3.row_dimensions[curr].height = 20
        for col_idx in range(1, len(a_row) + 1):
            cell = ws3.cell(row=curr, column=col_idx)
            cell.font = bold_font if col_idx in (1, 4) else regular_font
            cell.border = thin_border
            cell.alignment = left_align if col_idx in (2, 3) else center_align

    # Автоширина колонок для всех листов
    for ws in (ws1, ws2, ws3):
        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                val = str(cell.value or '')
                if len(val) > max_len and '\n' not in val:
                    max_len = len(val)
            ws.column_dimensions[col_letter].width = min(max(max_len + 3, 10), 50)

    ws1.column_dimensions["B"].width = 38
    ws1.column_dimensions["C"].width = 28
    ws2.column_dimensions["D"].width = 35
    ws2.column_dimensions["G"].width = 40
    ws2.column_dimensions["H"].width = 30
    ws2.column_dimensions["J"].width = 32
    ws2.column_dimensions["K"].width = 25
    ws2.column_dimensions["L"].width = 45

    ws3.column_dimensions["A"].width = 16
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 40
    ws3.column_dimensions["D"].width = 28
    ws3.column_dimensions["E"].width = 16
    ws3.column_dimensions["N"].width = 24

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


CATEGORY_FOLDERS = {
    "scopus_wos":       "01_Scopus_va_Web_of_Science",
    "dsc":              "02_DSc_Dissertatsiyalar",
    "phd":              "03_PhD_Dissertatsiyalar",
    "patent":           "04_Patentlar_IAP_FAP",
    "monography":       "05_Monografiyalar",
    "oak_uz":           "06_OzOAK_maqolalari",
    "oak_ru_if":        "07_Rossiya_va_Xorijiy_OAK",
    "thesis_uz":        "08_Respublika_konferensiya_tezislari",
    "thesis_foreign":   "09_Xorijiy_konferensiya_tezislari",
    "rationalizer":     "10_Ratsionalizatorlik_takliflari",
    "implementation":   "11_Amaliyotga_tadbiq_qilingan_ishlar",
    "conferences":      "12_Kafedra_otkazgan_anjumanlar",
    "contracts":        "13_Xojalik_shartnomalari",
    "grants":           "14_Grantlar",
}


def sanitize_filename(text: str, max_len: int = 50) -> str:
    """Очищает строку от недопустимых символов для имен файлов"""
    import re
    if not text:
        return "hujjat"
    clean = re.sub(r'[\\/*?:"<>|\n\r\t]', '_', str(text))
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean[:max_len]


async def generate_files_zip(bot, entries: list) -> tuple[io.BytesIO, int]:
    """
    Скачивает файлы из Telegram и упаковывает их в ZIP-архив,
    разложенный по папкам категорий и кафедр.
    Возвращает (io.BytesIO, count_files).
    """
    import zipfile
    from pathlib import Path

    buf = io.BytesIO()
    count = 0

    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for e in entries:
            file_id = e['file_id']
            if not file_id:
                continue

            try:
                # Скачиваем файл из серверов Telegram
                tg_file = await bot.get_file(file_id)
                f_stream = await bot.download_file(tg_file.file_path)
                
                if hasattr(f_stream, 'getvalue'):
                    content = f_stream.getvalue()
                elif hasattr(f_stream, 'read'):
                    if hasattr(f_stream, 'seek'):
                        f_stream.seek(0)
                    content = f_stream.read()
                else:
                    content = bytes(f_stream)

                ext = Path(tg_file.file_path).suffix or ".pdf"
                if not ext.startswith("."):
                    ext = "." + ext

                cat_folder = CATEGORY_FOLDERS.get(e['category'], e['category'])
                dept_id = e['dept_id']
                author = sanitize_filename(e['authors'], 25)
                title = sanitize_filename(e['title'], 35)

                # Пример имени: 01_Scopus_va_Web_of_Science/[Kaf_02]_Uzbekova_Digital_therapeutics_id12.pdf
                zip_path = f"{cat_folder}/[Kaf_{dept_id:02d}]_{author}_{title}_id{e['id']}{ext}"
                zf.writestr(zip_path, content)
                count += 1
            except Exception as ex:
                print(f"Failed to fetch file for entry #{e['id']}: {ex}")
                continue

    buf.seek(0)
    return buf, count


